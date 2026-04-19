from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from io import BytesIO
import json
from store import get_task



router = APIRouter(prefix="/download", tags=["download"])


@router.get("/{task_id}")
def download_result(task_id: str, format: str = "pdf"):
    task = get_task(task_id)
    if not task or not task.result:
        raise HTTPException(status_code=404, detail="Task result not found")        
    content = task.result.get("content", "")
    if format == "pdf":
        try:
            from fpdf import FPDF
            if isinstance(content, dict) and 'pages' in content:
                pdf = FPDF(orientation=content.get('orientation', 'portrait')[0].upper(), unit='pt', format=content.get('page_size', 'A4').upper())
                pdf.set_auto_page_break(auto=True, margin=72)
                for page in content.get('pages', []):
                    pdf.add_page()
                    if content.get('header'):
                        header = content['header']
                        pdf.set_font("Arial", size=8)
                        pdf.set_text_color(128, 128, 128)
                        if header.get('logo_placeholder'):
                            pdf.cell(0, 20, header['logo_placeholder'].encode('latin-1', 'replace').decode('latin-1'), ln=False)
                        if header.get('text'):
                            pdf.cell(0, 20, header['text'].encode('latin-1', 'replace').decode('latin-1'), ln=True, align='R')
                        pdf.line(72, pdf.get_y(), 595-72, pdf.get_y())
                        pdf.ln(10)
                        pdf.set_text_color(0, 0, 0)
                    for block in page.get('content', []):
                        block_type = block.get('type', '')
                        if block_type == 'title':
                            pdf.set_font("Arial", 'B', size=block.get('font_size', 24))
                            align = block.get('alignment', 'left')[0].upper()
                            pdf.cell(0, 30, block.get('text', '').encode('latin-1', 'replace').decode('latin-1'), ln=True, align=align)
                            pdf.ln(5)
                        elif block_type == 'paragraph':
                            pdf.set_font("Arial", size=block.get('font_size', 12))
                            pdf.multi_cell(0, block.get('font_size', 12) + 2, block.get('text', '').encode('latin-1', 'replace').decode('latin-1'))
                            pdf.ln(5)
                        elif block_type == 'table':
                            pdf.set_font("Arial", 'B', size=11)
                            headers = block.get('headers', [])
                            rows = block.get('rows', [])
                            col_width = (595 - 144) / len(headers) if headers else 100
                            pdf.set_fill_color(220, 220, 220)
                            for header in headers:
                                pdf.cell(col_width, 20, header.encode('latin-1', 'replace').decode('latin-1'), border=1, fill=True)
                            pdf.ln()
                            pdf.set_font("Arial", size=10)
                            for i, row in enumerate(rows):
                                if block.get('style') == 'striped' and i % 2 == 1:
                                    pdf.set_fill_color(240, 240, 240)
                                else:
                                    pdf.set_fill_color(255, 255, 255)
                                for cell in row:
                                    pdf.cell(col_width, 18, str(cell).encode('latin-1', 'replace').decode('latin-1'), border=1, fill=True)
                                pdf.ln()
                            pdf.ln(5)
                        elif block_type == 'list':
                            pdf.set_font("Arial", size=12)
                            items = block.get('items', [])
                            ordered = block.get('ordered', False)
                            for i, item in enumerate(items, 1):
                                prefix = f"{i}. " if ordered else "• "
                                pdf.multi_cell(0, 15, f"{prefix}{item}".encode('latin-1', 'replace').decode('latin-1'))
                            pdf.ln(5)
                        elif block_type == 'divider':
                            pdf.ln(5)
                            pdf.line(72, pdf.get_y(), 595-72, pdf.get_y())
                            pdf.ln(5)
                        elif block_type == 'code_block':
                            pdf.set_font("Courier", size=9)
                            pdf.set_fill_color(240, 240, 240)
                            code = block.get('code', '')
                            for line in code.split('\n'):
                                pdf.multi_cell(0, 12, line.encode('latin-1', 'replace').decode('latin-1'), fill=True)
                            pdf.ln(5)
                    if content.get('footer'):
                        footer = content['footer']
                        pdf.set_y(-72)
                        pdf.line(72, pdf.get_y(), 595-72, pdf.get_y())
                        pdf.ln(5)
                        pdf.set_font("Arial", size=8)
                        pdf.set_text_color(128, 128, 128)
                        if footer.get('show_date'):
                            pdf.cell(0, 15, footer.get('text', '').encode('latin-1', 'replace').decode('latin-1'), ln=False)
                        pdf.cell(0, 15, f"Page {page.get('page_number', 1)}", ln=True, align='R')
                        pdf.set_text_color(0, 0, 0)
                pdf_bytes = pdf.output(dest='S')
            elif isinstance(content, dict) and 'content' in content:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", 'B', size=16)
                title = content.get('title', 'Document')
                pdf.cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                pdf.ln(5)
                pdf.set_font("Arial", size=10)
                metadata = content.get('metadata', {})
                author = metadata.get('author', content.get('author', 'Unknown'))
                pdf.cell(0, 5, f"Author: {author}".encode('latin-1', 'replace').decode('latin-1'), ln=True)
                pdf.ln(5)
                for block in content.get('content', []):
                    block_type = block.get('type', '')
                    if block_type in ['heading', 'header']:
                        level = block.get('level', 2)
                        size = max(14 - level, 10)
                        pdf.set_font("Arial", 'B', size=size)
                        pdf.ln(3)
                        pdf.multi_cell(0, 6, block.get('text', '').encode('latin-1', 'replace').decode('latin-1'))
                        pdf.ln(2)
                    elif block_type == 'paragraph':
                        pdf.set_font("Arial", size=11)
                        pdf.multi_cell(0, 5, block.get('text', '').encode('latin-1', 'replace').decode('latin-1'))
                        pdf.ln(3)
                    elif block_type == 'code_block':
                        pdf.set_font("Courier", size=9)
                        pdf.set_fill_color(240, 240, 240)
                        code = block.get('code', '')
                        for line in code.split('\n'):
                            pdf.multi_cell(0, 4, line.encode('latin-1', 'replace').decode('latin-1'), fill=True)
                        pdf.ln(3)
                    elif block_type == 'list':
                        pdf.set_font("Arial", size=11)
                        items = block.get('items', [])
                        ordered = block.get('ordered', False)
                        for i, item in enumerate(items, 1):
                            prefix = f"{i}. " if ordered else "• "
                            pdf.multi_cell(0, 5, f"{prefix}{item}".encode('latin-1', 'replace').decode('latin-1'))
                        pdf.ln(3)
                citations = metadata.get('citations', content.get('citations', []))
                if citations:
                    pdf.ln(5)
                    pdf.set_font("Arial", 'B', size=12)
                    pdf.cell(0, 10, "References", ln=True)
                    pdf.set_font("Arial", size=9)
                    for i, cit in enumerate(citations, 1):
                        if isinstance(cit, dict):
                            cit_text = f"[{i}] {cit.get('text', '')} - {cit.get('url', '')}"
                        else:
                            cit_text = f"[{i}] {cit}"
                        pdf.multi_cell(0, 4, cit_text.encode('latin-1', 'replace').decode('latin-1'))
                        
                pdf_bytes = pdf.output(dest='S')
            else:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                text_content = str(content)
                for line in text_content.split("\n"):
                    clean_line = line.encode("latin-1", "replace").decode("latin-1")
                    pdf.multi_cell(0, 10, clean_line)
                pdf_bytes = pdf.output(dest='S')
            buffer = BytesIO(pdf_bytes)
            return StreamingResponse(
                buffer, 
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=result_{task_id}.pdf"}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="fpdf2 is not installed")
    elif format == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            doc = Document()
            if isinstance(content, dict):
                title = doc.add_heading(content.get('title', 'Document'), 0)
                metadata = content.get('metadata', {})
                author = metadata.get('author', content.get('author', 'Unknown'))
                p = doc.add_paragraph(f"Author: {author}")
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                for block in content.get('content', []):
                    block_type = block.get('type', '')
                    if block_type in ['heading', 'header']:
                        level = min(block.get('level', 2), 9)
                        doc.add_heading(block.get('text', ''), level)
                    elif block_type == 'paragraph':
                        doc.add_paragraph(block.get('text', ''))
                    elif block_type == 'code_block':
                        p = doc.add_paragraph(block.get('code', ''))
                        p.style = 'No Spacing'
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                    elif block_type == 'list':
                        items = block.get('items', [])
                        ordered = block.get('ordered', False)
                        style = 'List Number' if ordered else 'List Bullet'
                        for item in items:
                            doc.add_paragraph(item, style=style)
                citations = metadata.get('citations', content.get('citations', []))
                if citations:
                    doc.add_heading('References', 2)
                    for i, cit in enumerate(citations, 1):
                        if isinstance(cit, dict):
                            cit_text = f"[{i}] {cit.get('text', '')} - {cit.get('url', '')}"
                        else:
                            cit_text = f"[{i}] {cit}"
                        p = doc.add_paragraph(cit_text)
                        p.runs[0].font.size = Pt(9)
            else:
                doc.add_heading("Verification Result", 0)
                doc.add_paragraph(str(content))
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return StreamingResponse(
                buffer, 
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=result_{task_id}.docx"}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx is not installed")
    raise HTTPException(status_code=400, detail="Invalid format specified. Use 'pdf' or 'docx'.")